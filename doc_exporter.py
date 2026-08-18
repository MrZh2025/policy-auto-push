# -*- coding: utf-8 -*-
"""
党政机关标准公文排版与 Word 导出引擎（严格遵循 GB/T 9704-2012 国家标准）
特点：
1. 页面设置：标准 A4 型纸，上 37mm、下 35mm、左 28mm、右 26mm（符合装订线标准）
2. 字号与字体层级：
   - 公文标题：2号 方正小标宋简体（居中，行距32磅，不加标点）
   - 一级标题：3号 黑体（首行缩进2字符，如：一、政策速览与总体动向）
   - 二级标题：3号 楷体_GB2312（首行缩进2字符，如：（一）国务院及部委医药综合政策）
   - 三级标题：3号 方正仿宋简体加粗（首行缩进2字符，如：1. 医疗卫生强基工程中医药行动。）
   - 正文内容：3号 方正仿宋简体（首行缩进2字符，固定行距 28.5 磅）
3. 三线表：标准政务三线表（顶底粗线、栏目细线，表头黑体，单元格仿宋）
4. 页码规范：符合 GB/T 9704-2012 奇偶页不同两侧排布（单页码靠右空一字，双页码靠左空一字，格式为 — 1 —）
5. 结构纯净：干练总结、直击要点、无冗余落款与抄送版记
"""
import os
import sys
import io
import re
import logging
from datetime import datetime
from typing import List, Dict, Any, Optional

# 解决 Windows 控制台编码
if sys.platform == 'win32' and hasattr(sys.stdout, 'reconfigure'):
    try:
        sys.stdout.reconfigure(encoding='utf-8', errors='replace')
        sys.stderr.reconfigure(encoding='utf-8', errors='replace')
    except Exception:
        pass

from docx import Document
from docx.shared import Pt, Inches, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls

logger = logging.getLogger(__name__)

# ----------------- 基础字号对照表 (Pt) -----------------
FONT_SIZES = {
    '1号': Pt(26),
    '2号': Pt(22),
    '小2号': Pt(18),
    '3号': Pt(16),
    '小3号': Pt(15),
    '4号': Pt(14),
    '小4号': Pt(12),
    '5号': Pt(10.5)
}

# ----------------- 标准公文字体映射（中文字体 + Times New Roman 英文与数值） -----------------
FONT_NAMES = {
    'title': ('方正小标宋简体', 'FZXiaoBiaoSong-B05S', 'Times New Roman'),
    'body': ('方正仿宋简体', '仿宋_GB2312', '仿宋', 'Times New Roman'),
    'h1': ('黑体', 'SimHei', 'Times New Roman'),
    'h2': ('楷体_GB2312', 'KaiTi_GB2312', 'Times New Roman'),
    'h3': ('方正仿宋简体', '仿宋_GB2312', '仿宋', 'Times New Roman'),
    'table_header': ('黑体', 'SimHei', 'Times New Roman'),
    'table': ('方正仿宋简体', '仿宋_GB2312', '仿宋', 'Times New Roman'),
    'page_num': ('Times New Roman', 'SimSun', 'Times New Roman')
}

def set_run_font(run, font_type='body', size_pt=Pt(16), bold=False, italic=False, color_rgb=(0, 0, 0)):
    """设置文本中西文字体、字号与样式，严格绑定 eastAsia 为中文字体，ascii/hAnsi/cs 为 Times New Roman"""
    run.font.size = size_pt
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor(*color_rgb)
    
    font_tuple = FONT_NAMES.get(font_type, ('方正仿宋简体', '仿宋_GB2312', '仿宋', 'Times New Roman'))
    east_asia_font = font_tuple[0]
    ascii_font = 'Times New Roman'
    
    run.font.name = ascii_font
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
        rPr.append(rFonts)
    rFonts.set(qn('w:ascii'), ascii_font)
    rFonts.set(qn('w:hAnsi'), ascii_font)
    rFonts.set(qn('w:eastAsia'), east_asia_font)
    rFonts.set(qn('w:cs'), ascii_font)

def strip_markdown(text: str) -> str:
    """彻底消除文本中的 Markdown 标记（如 *、**、#、列表符号等）"""
    if not text:
        return ""
    s = str(text)
    s = re.sub(r'\[(.*?)\]\((.*?)\)', r'\1', s)
    s = re.sub(r'\*{3}(.*?)\*{3}', r'\1', s)
    s = re.sub(r'\*{2}(.*?)\*{2}', r'\1', s)
    s = re.sub(r'\*([^\*\n]+)\*', r'\1', s)
    s = re.sub(r'_{3}(.*?)_{3}', r'\1', s)
    s = re.sub(r'_{2}(.*?)_{2}', r'\1', s)
    s = re.sub(r'_([^_\n]+)_', r'\1', s)
    s = re.sub(r'^[\s\-\*\+•]+\s*', '', s)
    s = s.replace('*', '')
    s = re.sub(r'`([^`]+)`', r'\1', s)
    s = s.replace('`', '')
    s = re.sub(r'^#+\s*', '', s)
    return s.strip()

def set_para_spacing(p, line_spacing_rule=WD_LINE_SPACING.ONE_POINT_FIVE, space_before_pt=Pt(0), space_after_pt=Pt(0), align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line_indent_pt=Pt(24)):
    """设置段落 1.5 倍行距、首行缩进 2 字符 (小4号字为24pt) 与对齐"""
    p.paragraph_format.line_spacing_rule = line_spacing_rule
    p.paragraph_format.space_before = space_before_pt
    p.paragraph_format.space_after = space_after_pt
    p.paragraph_format.alignment = align
    if first_line_indent_pt is not None and first_line_indent_pt >= 0:
        p.paragraph_format.first_line_indent = first_line_indent_pt

def set_cell_border(cell, **kwargs):
    """设置单元格边框"""
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)
            for key, val in edge_data.items():
                element.set(qn('w:{}'.format(key)), str(val))

def add_side_page_number(doc):
    """
    添加符合国家标准 GB/T 9704-2012 规定的奇偶页不同两侧阿拉伯数字页码：
    - 单页码（奇数页）：居右空一字，格式为 — 1 —
    - 双页码（偶数页）：居左空一字，格式为 — 2 —
    """
    doc.settings.odd_and_even_pages_header_footer = True
    settings_elm = doc.settings._element
    even_odd = settings_elm.find(qn('w:evenAndOddHeaders'))
    if even_odd is not None:
        even_odd.set(qn('w:val'), '1')
    else:
        new_even_odd = parse_xml(r'<w:evenAndOddHeaders %s w:val="1"/>' % nsdecls('w'))
        settings_elm.append(new_even_odd)
    
    def setup_footer_content(footer, is_odd_page=True):
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if is_odd_page else WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.first_line_indent = Pt(0)
        p.text = ""
        
        # 偶数页（双页码）居左空一字
        if not is_odd_page:
            r_space_left = p.add_run("　")
            set_run_font(r_space_left, font_type='page_num', size_pt=FONT_SIZES['4号'])
            
        r_pre = p.add_run("— ")
        set_run_font(r_pre, font_type='page_num', size_pt=FONT_SIZES['4号'])
        
        # 动态 PAGE 域
        r_begin = p.add_run()
        r_begin._r.append(parse_xml(r'<w:fldChar %s w:fldCharType="begin"/>' % nsdecls('w')))
        r_instr = p.add_run()
        r_instr._r.append(parse_xml(r'<w:instrText %s xml:space="preserve"> PAGE </w:instrText>' % nsdecls('w')))
        r_sep = p.add_run()
        r_sep._r.append(parse_xml(r'<w:fldChar %s w:fldCharType="separate"/>' % nsdecls('w')))
        r_num = p.add_run("1" if is_odd_page else "2")
        set_run_font(r_num, font_type='page_num', size_pt=FONT_SIZES['4号'])
        r_end = p.add_run()
        r_end._r.append(parse_xml(r'<w:fldChar %s w:fldCharType="end"/>' % nsdecls('w')))
        
        r_suf = p.add_run(" —")
        set_run_font(r_suf, font_type='page_num', size_pt=FONT_SIZES['4号'])
        
        # 奇数页（单页码）居右空一字
        if is_odd_page:
            r_space_right = p.add_run("　")
            set_run_font(r_space_right, font_type='page_num', size_pt=FONT_SIZES['4号'])

    for section in doc.sections:
        setup_footer_content(section.footer, is_odd_page=True)
        setup_footer_content(section.even_page_footer, is_odd_page=False)

class PolicyDocExporter:
    """政策公文标准 Word 导出器"""

    @staticmethod
    def get_desktop_path() -> str:
        desktop = os.path.join(os.path.expanduser("~"), "Desktop")
        if os.path.exists(desktop):
            return desktop
        onedrive_desktop = os.path.join(os.path.expanduser("~"), "OneDrive", "Desktop")
        if os.path.exists(onedrive_desktop):
            return onedrive_desktop
        return os.path.expanduser("~")

    @classmethod
    def export(cls, policies: List[Dict[str, Any]], custom_filename: Optional[str] = None) -> List[str]:
        if not policies:
            logger.info("暂无政策，跳过 Word 导出。")
            return []

        doc = Document()

        # 1. 页面设置：严格遵循 GB/T 9704-2012 规定 (A4, 上37mm, 下35mm, 左28mm, 右26mm)
        section = doc.sections[0]
        section.page_width = Mm(210)
        section.page_height = Mm(297)
        section.top_margin = Mm(37)
        section.bottom_margin = Mm(35)
        section.left_margin = Mm(28)
        section.right_margin = Mm(26)

        today_str = datetime.now().strftime("%Y年%m月%d日")
        date_tag = datetime.now().strftime("%Y%m%d")

        # 2. 公文大标题：2号 方正小标宋简体，居中排布，1.5 倍行距，末尾不加标点
        p_title = doc.add_paragraph()
        set_para_spacing(p_title, line_spacing_rule=WD_LINE_SPACING.ONE_POINT_FIVE, space_before_pt=Pt(6), space_after_pt=Pt(12), first_line_indent_pt=Pt(0), align=WD_ALIGN_PARAGRAPH.CENTER)
        r_title = p_title.add_run("四川生物医药产业集团创新事业部政策信息简报")
        set_run_font(r_title, font_type='title', size_pt=FONT_SIZES['2号'], bold=True)

        # 3. 导语段落：小4号 方正仿宋简体（英文和数字 Times New Roman），首行缩进 2 字符 (24pt)，1.5 倍行距
        p_lead = doc.add_paragraph()
        set_para_spacing(p_lead, line_spacing_rule=WD_LINE_SPACING.ONE_POINT_FIVE, space_before_pt=Pt(0), space_after_pt=Pt(4), first_line_indent_pt=Pt(24), align=WD_ALIGN_PARAGRAPH.JUSTIFY)
        r_lead = p_lead.add_run(f"为及时研判行业监管动向与政策红利，现将截至{today_str}本周最新发布的医药产业重点政策及文件摘要汇总如下：")
        set_run_font(r_lead, font_type='body', size_pt=FONT_SIZES['小4号'])

        # 4. 一级标题：一、本周重点政策速览清单（小4号 黑体，首行缩进 2 字符，1.5 倍行距，末尾不加标点）
        p_h1 = doc.add_paragraph()
        set_para_spacing(p_h1, line_spacing_rule=WD_LINE_SPACING.ONE_POINT_FIVE, space_before_pt=Pt(6), space_after_pt=Pt(3), first_line_indent_pt=Pt(24), align=WD_ALIGN_PARAGRAPH.JUSTIFY)
        r_h1 = p_h1.add_run("一、本周重点政策速览清单")
        set_run_font(r_h1, font_type='h1', size_pt=FONT_SIZES['小4号'])

        # 5. 标准公文三线表（顶底线粗 1.5pt，栏目线 0.75pt，无竖线）
        table = doc.add_table(rows=len(policies) + 1, cols=4)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False

        # 设置列宽
        col_widths = [Mm(14), Mm(80), Mm(38), Mm(24)]
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = w

        # 表头 (小4号 黑体居中)
        headers = ["序号", "政策文件名称", "发布机关", "发布日期"]
        for i, h_text in enumerate(headers):
            cell = table.cell(0, i)
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.first_line_indent = Pt(0)
            r = p.add_run(h_text)
            set_run_font(r, font_type='table_header', size_pt=FONT_SIZES['小4号'], bold=True)
            # 顶线 1.5pt，栏目线 0.75pt
            set_cell_border(cell, top={'val': 'single', 'sz': '12', 'color': '000000'},
                                  bottom={'val': 'single', 'sz': '6', 'color': '000000'},
                                  left={'val': 'none'}, right={'val': 'none'})

        # 数据行 (方正仿宋简体，小4号)
        for idx, item in enumerate(policies, 1):
            row_cells = table.rows[idx].cells
            row_data = [
                str(idx),
                strip_markdown(item.get("title", "")),
                strip_markdown(item.get("source", "")),
                strip_markdown(item.get("pub_date", "") or "近期")
            ]
            is_last_row = (idx == len(policies))

            for col_idx, text_val in enumerate(row_data):
                cell = row_cells[col_idx]
                cell.text = ""
                p = cell.paragraphs[0]
                p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                p.paragraph_format.space_before = Pt(3)
                p.paragraph_format.space_after = Pt(3)
                p.paragraph_format.first_line_indent = Pt(0)
                if col_idx in [0, 3]:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

                r = p.add_run(text_val)
                set_run_font(r, font_type='table', size_pt=FONT_SIZES['小4号'])

                # 底边线设置（最后一行底线 1.5pt，中间无横线）
                if is_last_row:
                    set_cell_border(cell, bottom={'val': 'single', 'sz': '12', 'color': '000000'},
                                          top={'val': 'none'}, left={'val': 'none'}, right={'val': 'none'})
                else:
                    set_cell_border(cell, bottom={'val': 'none'}, top={'val': 'none'},
                                          left={'val': 'none'}, right={'val': 'none'})

        # 6. 一级标题：二、本周重点政策要点与文件摘要（小4号 黑体，首行缩进 2 字符，1.5 倍行距）
        p_h2 = doc.add_paragraph()
        set_para_spacing(p_h2, line_spacing_rule=WD_LINE_SPACING.ONE_POINT_FIVE, space_before_pt=Pt(8), space_after_pt=Pt(3), first_line_indent_pt=Pt(24), align=WD_ALIGN_PARAGRAPH.JUSTIFY)
        r_h2 = p_h2.add_run("二、本周重点政策要点与文件摘要")
        set_run_font(r_h2, font_type='h1', size_pt=FONT_SIZES['小4号'])

        # 7. 政策要点逐条排版（遵循规范层级：小4号 标目标题 + 紧凑要点正文，1.5 倍行距，首行缩进 2 字符）
        for idx, item in enumerate(policies, 1):
            title = strip_markdown(item.get('title', ''))
            dept = strip_markdown(item.get("source", "国家部委"))
            date_str = strip_markdown(item.get("pub_date", "") or "近期")
            summary = strip_markdown(item.get("summary", "") or title)
            url = item.get("url", "")

            # 条目标题段落（小4号 方正仿宋加粗，首行缩进 2 字符，1.5 倍行距）
            p_item = doc.add_paragraph()
            set_para_spacing(p_item, line_spacing_rule=WD_LINE_SPACING.ONE_POINT_FIVE, space_before_pt=Pt(4), space_after_pt=Pt(0), first_line_indent_pt=Pt(24), align=WD_ALIGN_PARAGRAPH.JUSTIFY)
            r_ititle = p_item.add_run(f"{idx}. 《${title}》（发布机关：${dept}，发布日期：${date_str}）".replace('$', ''))
            set_run_font(r_ititle, font_type='h3', size_pt=FONT_SIZES['小4号'], bold=True)

            # 内容与要点正文（小4号 方正仿宋简体，首行缩进 2 字符，1.5 倍行距，附带官方链接）
            p_desc = doc.add_paragraph()
            set_para_spacing(p_desc, line_spacing_rule=WD_LINE_SPACING.ONE_POINT_FIVE, space_before_pt=Pt(0), space_after_pt=Pt(3), first_line_indent_pt=Pt(24), align=WD_ALIGN_PARAGRAPH.JUSTIFY)
            r_desc = p_desc.add_run(f"文件主要内容与核心要点：{summary}")
            set_run_font(r_desc, font_type='body', size_pt=FONT_SIZES['小4号'])

            if url and url != '#':
                r_link_label = p_desc.add_run("（官方原文直达：")
                set_run_font(r_link_label, font_type='body', size_pt=FONT_SIZES['小4号'])
                r_link = p_desc.add_run(url)
                set_run_font(r_link, font_type='body', size_pt=FONT_SIZES['小4号'], color_rgb=(0, 72, 134))
                r_link.font.underline = True
                r_link_close = p_desc.add_run("）")
                set_run_font(r_link_close, font_type='body', size_pt=FONT_SIZES['小4号'])

        # 8. 奇偶页不同两侧页码（单页居右空一字，双页居左空一字，格式为 — 1 —）
        add_side_page_number(doc)

        # 10. 保存文件
        filename = custom_filename or f"四川生物医药产业集团创新事业部政策信息简报_{date_tag}.docx"
        saved_paths = []

        # 9.1 保存到 Windows 桌面
        desktop_dir = cls.get_desktop_path()
        desktop_file = os.path.join(desktop_dir, filename)
        try:
            doc.save(desktop_file)
            saved_paths.append(desktop_file)
            logger.info(f"📄 [公文Word已保存到桌面] {desktop_file}")
        except Exception as e:
            logger.warning(f"保存到桌面失败: {e}")

        # 9.2 在项目内 reports/ 目录保存归档
        project_reports_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "reports")
        os.makedirs(project_reports_dir, exist_ok=True)
        report_file = os.path.join(project_reports_dir, filename)
        try:
            doc.save(report_file)
            saved_paths.append(report_file)
            logger.info(f"📁 [公文Word已归档到项目目录] {report_file}")
        except Exception as e:
            logger.warning(f"保存到 reports/ 失败: {e}")

        return saved_paths


if __name__ == "__main__":
    logging.basicConfig(level=logging.INFO, format="%(asctime)s [%(levelname)s] %(message)s")
    from database import PolicyDatabase
    db = PolicyDatabase()
    policies = db.get_unpushed_policies(limit=6)
    if not policies:
        with db._get_connection() as conn:
            cursor = conn.cursor()
            cursor.execute("SELECT * FROM policies ORDER BY id DESC LIMIT 6")
            policies = [dict(row) for row in cursor.fetchall()]

    paths = PolicyDocExporter.export(policies)
    print("✅ 公文格式 Word 导出完成:")
    for p in paths:
        print(" ->", p)
